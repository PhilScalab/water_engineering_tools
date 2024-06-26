import openpyxl
from openpyxl.utils.dataframe import dataframe_to_rows
from openpyxl.drawing.image import Image
import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import io
from io import StringIO
import base64
import docx
import matplotlib.dates as mdates
import scipy.stats as stats
from scipy.stats import norm, lognorm, gumbel_r,pearson3, gamma, genextreme
#from scipy.stats import stats, norm, lognorm, pearson3, gamma, gumbel_r, genextreme
from openpyxl import Workbook
from openpyxl.utils.dataframe import dataframe_to_rows
from openpyxl.drawing.image import Image
import tempfile
import os
import shutil
import zipfile
from datetime import datetime, timedelta
import math
import requests
from io import BytesIO
import pydeck as pdk


# Functions for Frequency Analysis


def log_pearson3(x, loc, scale, skew):
    return pearson3.pdf(x, skew, loc, scale)


def fit_distribution(distr, data):
    params = distr.fit(data)
    log_likelihood = np.sum(np.log(distr.pdf(data, *params)))
    k = len(params)
    n = len(data)

    aic = 2 * k - 2 * log_likelihood
    bic = k * np.log(n) - 2 * log_likelihood

    return aic, bic, params


distributions = {
    'Normal': norm,
    'Lognormal': lognorm,
    'Pearson Type 3': pearson3,
    'Gamma': gamma,
    'Gumbel': gumbel_r,
    'GEV': genextreme,
}

def to_excel(df):
    output = BytesIO()
    writer = pd.ExcelWriter(output, engine='openpyxl')
    df.to_excel(writer, index=False, sheet_name='Sheet1')
    writer.save()
    processed_data = output.getvalue()
    return processed_data
    
def generate_word_document(max_flow, aic_bic_params, best_aic_distr, best_bic_distr):
    # Create a Word document
    doc = docx.Document()
    doc.add_heading('Frequency Analysis of Maximum Flow in Rivers', 0)

    doc.add_heading('Best Distribution based on AIC and BIC:', level=1)
    doc.add_paragraph(
        f"Best distribution based on AIC: {best_aic_distr} (AIC: {aic_bic_params[best_aic_distr]['AIC']})")
    doc.add_paragraph(
        f"Best distribution based on BIC: {best_bic_distr} (BIC: {aic_bic_params[best_bic_distr]['BIC']})")

    doc.add_heading('AIC and BIC for each distribution:', level=1)
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Distribution'
    hdr_cells[1].text = 'AIC'
    hdr_cells[2].text = 'BIC'

    for name, info in aic_bic_params.items():
        row_cells = table.add_row().cells
        row_cells[0].text = name
        row_cells[1].text = str(info['AIC'])
        row_cells[2].text = str(info['BIC'])

    doc.add_heading('Individual Distribution Plots:', level=1)
    for name in aic_bic_params.keys():
        doc.add_picture(f'{name}_distribution.png',
                        width=docx.shared.Inches(6))

    return doc
    
def download_ndbc_full_history(station_id, start_year, end_year):
    """
    Download full historical wave data for a given station across a range of years.

    Parameters:
    - station_id (str): The ID of the NDBC station.
    - start_year (int): The starting year for the data.
    - end_year (int): The ending year for the data.

    Returns:
    - DataFrame: A pandas DataFrame containing the downloaded data.
    """

    all_data = []

    for year in range(start_year, end_year + 1):
        # Base URL for historical data
        base_url = f"https://www.ndbc.noaa.gov/view_text_file.php?filename={station_id}h{year}.txt.gz&dir=data/historical/stdmet/"

        # Download the data
        response = requests.get(base_url)
        if response.status_code == 200:
            # Parse the data
            content = response.content.decode('utf-8')
            data = [line.split() for line in content.split('\n') if line and not line.startswith('#')]
            all_data.extend(data)
        else:
            st.write(f"Data not found for station {station_id} in year {year}.")

    # Create a DataFrame
    df = pd.DataFrame(all_data, columns=["YY", "MM", "DD", "hh", "mm", "WDIR", "WSPD", "GST", "WVHT", "DPD", "APD", "MWD", "PRES", "ATMP", "WTMP", "DEWP", "VIS", "TIDE"])
    
    return df
    
# def plot_wave_height(df, station_id):
#     """
#     Plot the Wave Height (WVHT) from the DataFrame.

#     Parameters:
#     - df (DataFrame): The DataFrame containing the wave data.
#     - station_id (str): The ID of the NDBC station.
#     """
#     plt.figure(figsize=(10, 6))
#     plt.plot(df['WVHT'], label='Wave Height (WVHT)')
#     plt.title(f'Wave Height (WVHT) at Station {station_id}')
#     plt.xlabel('Record Number')
#     plt.ylabel('Wave Height (meters)')
#     plt.legend()
#     st.pyplot(plt)
def plot_wave_height(df, title):
    """
    Plot the Wave Height (WVHT) from the DataFrame.

    Parameters:
    - df (DataFrame): The DataFrame containing the wave data.
    - title (str): The title for the plot.
    """
    plt.figure(figsize=(10, 6))
    plt.plot(df['WVHT'], label='Wave Height (WVHT)')
    plt.title(title)
    plt.xlabel('Date')
    plt.ylabel('Wave Height (meters)')
    plt.legend()
    st.pyplot(plt)

def download_link(document, filename):
    with io.BytesIO() as buffer:
        document.save(buffer)
        buffer.seek(0)
        file = base64.b64encode(buffer.read()).decode('utf-8')
    return f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{file}" download="{filename}">Download Word document</a>'

# Function to create a download link for the generated Excel file

def download_excel_link(excel_file, filename):
    with io.BytesIO() as buffer:
        excel_file.save(buffer)
        buffer.seek(0)
        file = base64.b64encode(buffer.read()).decode('utf-8')
    return f'<a href="data:application/vnd.openxmlformats-officedocument.spreadsheetml.sheet;base64,{file}" download="{filename}">Download Excel file</a>'


def generate_hydrographs_and_tables(daily_flow_data, sep_day, sep_month, spring_volume_period, fall_volume_period):
    unique_years = daily_flow_data.index.year.unique()

    # Create a new Excel workbook
    wb = Workbook()
    ws1 = wb.active
    ws1.title = "Hydrographs"

    max_spring_df = pd.DataFrame(
        columns=["Year", "Max Flow Spring", "Max Flow Date"])
    min_spring_df = pd.DataFrame(
        columns=["Year", "Min Flow Spring", "Min Flow Date"])
    max_fall_df = pd.DataFrame(
        columns=["Year", "Max Flow Fall", "Max Flow Date"])
    min_fall_df = pd.DataFrame(
        columns=["Year", "Min Flow Fall", "Min Flow Date"])
    period_df = pd.DataFrame(columns=["Year", "Spring Period", "Fall Period"])

    for year in unique_years:
        yearly_data = daily_flow_data[daily_flow_data.index.year == year]

        # Spring and Fall data
        spring_data = yearly_data.loc[yearly_data.index < yearly_data.index[0].replace(
            month=sep_month, day=sep_day)]
        fall_data = yearly_data.loc[yearly_data.index >= yearly_data.index[0].replace(
            month=sep_month, day=sep_day)]

        # Compute statistics
        spring_max_flow = spring_data['Flow'].max()
        spring_min_flow = spring_data['Flow'].min()
        fall_max_flow = fall_data['Flow'].max()
        fall_min_flow = fall_data['Flow'].min()
        spring_max_date = spring_data['Flow'].idxmax(
        ) if not spring_data.empty else None
        spring_min_date = spring_data['Flow'].idxmin(
        ) if not spring_data.empty else None
        fall_max_date = fall_data['Flow'].idxmax(
        ) if not fall_data.empty else None
        fall_min_date = fall_data['Flow'].idxmin(
        ) if not fall_data.empty else None

        # Add data to summary tables
        max_spring_df = max_spring_df.append(
            {"Year": year, "Max Flow Spring": spring_max_flow, "Max Flow Date": spring_max_date.strftime('%d-%m') if spring_max_date is not None else None}, ignore_index=True)
        min_spring_df = min_spring_df.append(
            {"Year": year, "Min Flow Spring": spring_min_flow, "Min Flow Date": spring_min_date.strftime('%d-%m')if spring_min_date is not None else None}, ignore_index=True)
        max_fall_df = max_fall_df.append({"Year": year, "Max Flow Fall": fall_max_flow,
                                          "Max Flow Date": fall_max_date.strftime('%d-%m')if fall_max_date is not None else None}, ignore_index=True)
        min_fall_df = min_fall_df.append({"Year": year, "Min Flow Fall": fall_min_flow,
                                          "Min Flow Date": fall_min_date.strftime('%d-%m')if fall_min_date is not None else None}, ignore_index=True)

        # Plot hydrograph
        fig, ax = plt.subplots(figsize=(10, 6))
        ax.plot(yearly_data.index, yearly_data['Flow'], label="Flow")

        # Add max and min points
        if spring_max_date is not None and spring_max_flow is not None:
            ax.plot(spring_max_date, spring_max_flow,
                    'ro', label="Max (Spring)")
        if spring_min_date is not None and spring_min_flow is not None:
            ax.plot(spring_min_date, spring_min_flow,
                    'go', label="Min (Spring)")
        if fall_max_date is not None and fall_max_flow is not None:
            ax.plot(fall_max_date, fall_max_flow, 'ro', label="Max (Fall)")
        if fall_min_date is not None and fall_min_flow is not None:
            ax.plot(fall_min_date, fall_min_flow, 'go', label="Min (Fall)")

        # Add separation date and spring/fall volume periods
        separation_date = yearly_data.index[0].replace(
            month=sep_month, day=sep_day)
        ax.axvline(separation_date, linestyle='--',
                   color='k', label="Separation Date")

        spring_rolling_sum = spring_data['Flow'].rolling(
            spring_volume_period).sum()
        if not spring_rolling_sum.empty:
            spring_volume_start = spring_rolling_sum.idxmax(
            ) - pd.Timedelta(days=spring_volume_period - 1)
            spring_volume_end = spring_rolling_sum.idxmax()
        else:
            spring_volume_start = None
            spring_volume_end = None
        ax.axvspan(spring_volume_start, spring_volume_end,
                   color='r', alpha=0.3, label="Spring Volume Period")

        fall_rolling_sum = fall_data['Flow'].rolling(fall_volume_period).sum()
        if not fall_rolling_sum.empty:
            fall_volume_start = fall_rolling_sum.idxmax(
            ) - pd.Timedelta(days=fall_volume_period - 1)
            fall_volume_end = fall_rolling_sum.idxmax()
        else:
            fall_volume_start = None
            fall_volume_end = None
        ax.axvspan(fall_volume_start, fall_volume_end, color='g',
                   alpha=0.3, label="Fall Volume Period")

        ax.set_title(f"Hydrograph {year}")
        ax.set_xlabel("Date")
        ax.set_ylabel("Flow")
        ax.legend(loc="best")
        ax.xaxis.set_major_formatter(mdates.DateFormatter("%b"))

        # Save the figure to a temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False)
        fig.savefig(temp_file.name, format="png", dpi=300, bbox_inches="tight")
        plt.close(fig)

        # Add the image to the Excel workbook
        img = Image(temp_file.name)
        img.width = img.width // 4
        img.height = img.height // 4
        ws1.column_dimensions["A"].width = img.width // 6
        ws1.row_dimensions[year - unique_years.min()].height = img.height
        ws1.add_image(img, f"A{year - unique_years.min() + 1}")

        # Delete the temporary file
        # os.unlink(temp_file.name)

        # Add data to periods table
        period_df = period_df.append({"Year": year, "Spring Period": f"{spring_volume_start.strftime('%d-%m') if spring_volume_start is not None else None} - {spring_volume_end.strftime('%d-%m') if spring_volume_end is not None else None}",
                                      "Fall Period": f"{fall_volume_start.strftime('%d-%m') if fall_volume_start is not None else None} - {fall_volume_end.strftime('%d-%m') if fall_volume_end is not None else None}"}, ignore_index=True)

    # Create remaining sheets in the Excel workbook
    for sheet_name, df in zip(["Max Spring", "Min Spring", "Max Fall", "Min Fall", "Periods"], [max_spring_df, min_spring_df, max_fall_df, min_fall_df, period_df]):
        ws = wb.create_sheet(sheet_name)
        for r in dataframe_to_rows(df, index=False, header=True):
            ws.append(r)

    # # Delete temporary image files
    # for year in range(min_year, max_year + 1):
    #     try:
    #         os.remove(f"temp_image_{year}.png")
    #     except FileNotFoundError:
    #         pass
    return wb


def download_link(workbook, filename):
    with io.BytesIO() as buffer:
        workbook.save(buffer)
        buffer.seek(0)
        file = base64.b64encode(buffer.read()).decode('utf-8')
    return f'<a href="data:application/vnd.openxmlformats-officedocument.spreadsheetml.sheet;base64,{file}" download="{filename}">Download Excel file</a>'
    
# Function to process the uploaded file
def process_file(uploaded_file):
    df = pd.read_csv(uploaded_file)

    # Check if the Year, Month, Day columns are present
    if {'Year', 'Month', 'Day'}.issubset(df.columns):
        # Construct the Date column from Year, Month, Day
        df['Date'] = pd.to_datetime(df[['Year', 'Month', 'Day']])
        return df
    else:
        st.error("Required columns 'Year', 'Month', and 'Day' are not present in the CSV file.")
        return None
#Plot normal distribution
def create_plot_normal(data):
    mu, std = stats.norm.fit(data)

    def create_plots_normal():
        fig, axs = plt.subplots(2, 2, figsize=(10, 8))

        # Histogram and theoretical density
        x = np.linspace(min(data), max(data), 100)
        axs[0, 0].hist(data, bins=30, density=True, alpha=0.5)
        axs[0, 0].plot(x, stats.norm.pdf(x, mu, std), 'k', linewidth=2)
        axs[0, 0].set_title('Histogram and Theoretical Density')

        # Q-Q plot
        stats.probplot(data, dist="norm", plot=axs[0, 1])
        axs[0, 1].set_title('Q-Q Plot')

        # CDF
        sorted_data = np.sort(data)
        empirical_cdf = np.arange(1, len(sorted_data)+1) / len(sorted_data)
        axs[1, 0].plot(sorted_data, empirical_cdf, marker='.', linestyle='none')
        axs[1, 0].plot(x, stats.norm.cdf(x, mu, std), 'k', linewidth=2)
        axs[1, 0].set_title('Empirical and Theoretical CDF')

        # P-P plot
        theoretical_cdf = stats.norm.cdf(sorted_data, mu, std)
        axs[1, 1].plot(empirical_cdf, theoretical_cdf, marker='.', linestyle='none')
        axs[1, 1].plot([0, 1], [0, 1], 'k--')
        axs[1, 1].set_title('P-P Plot')

        plt.tight_layout()
        plt.show()
# Function to calculate AIC and BIC for a given distribution and data
def calculate_aic_bic(data, distribution, *params):
    # Calculate the log likelihood
    log_likelihood = np.sum(distribution.logpdf(data, *params))
    
    # Calculate the number of parameters (including the location and scale parameters)
    k = len(params) + 2  # adding 2 for loc and scale
    
    # Calculate AIC and BIC
    aic = 2 * k - 2 * log_likelihood
    bic = k * np.log(len(data)) - 2 * log_likelihood
    
    return aic, bic

# Function to fit distribution and calculate AIC and BIC
def fit_and_calculate_criteria(data, distribution, name, is_log_transformed=False):
    if is_log_transformed:
        # For log-transformed distributions, transform the data before fitting
        data_to_fit = np.log(data)
    else:
        data_to_fit = data

    # Fit the distribution to the data
    params = distribution.fit(data_to_fit)

    # Calculate AIC and BIC
    aic, bic = calculate_aic_bic(data_to_fit, distribution, *params)
    
    return {
        "Distribution": name,
        "AIC": aic,
        "BIC": bic,
        "Parameters": params
    }

    create_plots_normal()
    
def process_hdw_file(file_path, node):
    # Column names for the second type of lines (assuming type 2 has the node data)
    columns_type2 = ['i', 'jamfem', 'thifems', 'thifemf', 'qx', 'qy', 'eta', 
                     'eta1', 'detax', 'detay', 'uwat', 'vwat', 'htw', 'tw', 
                     'cv', 'han', 'hun']

    # Read the file and extract the time and the row for the specified node
    with open(file_path, 'r') as file:
        first_line = next(file).strip().split()
        time_value = first_line[0]  # Extracting time from the first row, first column

        for line in file:
            split_line = line.strip().split()
            if len(split_line) == len(columns_type2) and split_line[0] == str(node):
                row = [time_value] + split_line  # Prepending time to the row data
                return pd.DataFrame([row], columns=['time'] + columns_type2)
    
    # Return an empty DataFrame if the node is not found
    return pd.DataFrame(columns=['time'] + columns_type2)

# Page configuration
st.set_page_config(page_title="Water Engineering Tools", layout="wide")

# Main menu
menu = ["Home", "Hydrograph Producer","Ice Analysis - En","Survey Planner","CrissPy","Analyse de la glace - Fr", "Peak Flow Comparison",
        "Camera Viewer", "Frequency Analysis","EC Daily Data Analysis","Water level CEHQ","NDBC Historical Data Download","Frequency Analysis v2"]
choice = st.sidebar.selectbox("Menu", menu)

#CrissPy
if choice == "CrissPy":
    # Title of the app
    st.title("🥨 CrissPy")

    if 'combined_data' not in st.session_state:
        st.session_state['combined_data'] = pd.DataFrame()

    # File upload
    uploaded_file = st.file_uploader("Upload a zip file containing HDW files", type="zip")

    if uploaded_file is not None:
        # Node selection
        node = st.number_input("Select the node number", min_value=1)

        # Process button
        process_button = st.button("Process Data")

        if process_button:
            with zipfile.ZipFile(uploaded_file, 'r') as zip_ref:
                # Extract files to a temporary directory
                temp_dir = "temp_extract"
                zip_ref.extractall(temp_dir)

                # Process each .hdw file and concatenate the data for the selected node
                combined_data = pd.DataFrame()
                for filename in os.listdir(temp_dir):
                    if filename.endswith('.hdw'):
                        file_path = os.path.join(temp_dir, filename)
                        df = process_hdw_file(file_path, node)
                        combined_data = pd.concat([combined_data, df])

                # Cleanup temporary files
                for filename in os.listdir(temp_dir):
                    os.remove(os.path.join(temp_dir, filename))
                os.rmdir(temp_dir)

                # Store the combined data in session state
                st.session_state['combined_data'] = combined_data

    if not st.session_state['combined_data'].empty:
        # Convert time to numeric and sort data
        st.session_state['combined_data']['time'] = pd.to_numeric(st.session_state['combined_data']['time'])
        st.session_state['combined_data'].sort_values(by='time', inplace=True)

        # Column selection for plotting
        column = st.selectbox("Select a column for plotting", st.session_state['combined_data'].columns[1:])  # Exclude the time column

        # Process button
        plot_button = st.button("Plot Data")

        if plot_button:
            # Plotting
            fig, ax = plt.subplots(figsize=(10, 6))
            ax.plot(st.session_state['combined_data']['time'], st.session_state['combined_data'][column], marker='', color='blue', linewidth=2)
            ax.set_title(f"{column} over Time for Node {node}", fontsize=16)
            ax.set_xlabel("Time (hours)", fontsize=14)
            ax.set_ylabel(column, fontsize=14)
            ax.tick_params(axis='x', labelsize=12)
            ax.tick_params(axis='y', labelsize=12)
            ax.set_xlim([st.session_state['combined_data']['time'].min(), st.session_state['combined_data']['time'].max()])
            ax.set_ylim([st.session_state['combined_data'][column].min(), st.session_state['combined_data'][column].max()])
            ax.yaxis.set_major_locator(plt.MaxNLocator(10))
    
            st.pyplot(fig)
        # Generate a download button for the DataFrame
        # Download button for the DataFrame
        df_excel = to_excel(st.session_state['combined_data'])
        st.download_button(
            label="Download Excel file",
            data=df_excel,
            file_name="node_data.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
    else:
        st.write("Upload a file and process the data to view results.")


#"EWS-GS : Early warning system - Gauge Prediction"
if choice == "Survey Planner":
    # Title of the app
    st.title("🌧️ Survey Planner")
    
    # Function to perform regression calculation
    def calculate_predicted_flow(rain_data, weights):
        # Assuming a simple linear regression model: flow = sum(rain_data * weights)
        return np.sum(np.array(rain_data) * np.array(weights))
    
    # Creating two columns for Québec and Lévis
    col1, col2 = st.beta_columns(2)
    
    # Sector: Québec
    with col1:
        st.header("Québec")
        rain_today_qc = st.slider("Rain Today in Québec (mm)", 0, 100, 25, key="rain_today_qc")
        rain_tomorrow_qc = st.slider("Predicted Rain Tomorrow in Québec (mm)", 0, 100, 25, key="rain_tomorrow_qc")
        rain_day_after_qc = st.slider("Predicted Rain 2 Days from Now in Québec (mm)", 0, 100, 25, key="rain_day_after_qc")
    
        # Regression weights for Québec
        weights_qc = [0.29949278, 0.29876919, 0.22330731]
    
        # Perform calculations for Québec
        rain_data_qc = [rain_today_qc, rain_tomorrow_qc, rain_day_after_qc]
        predicted_flow_qc = calculate_predicted_flow(rain_data_qc, weights_qc)
        st.write(f"Predicted Flow in Québec: {predicted_flow_qc} m³/s")
    
    # Sector: Lévis
    with col2:
        st.header("Lévis")
        rain_today_levis = st.slider("Rain Today in Lévis (mm)", 0, 100, 25, key="rain_today_levis")
        rain_tomorrow_levis = st.slider("Predicted Rain Tomorrow in Lévis (mm)", 0, 100, 25, key="rain_tomorrow_levis")
        rain_day_after_levis = st.slider("Predicted Rain 2 Days from Now in Lévis (mm)", 0, 100, 25, key="rain_day_after_levis")
    
        # Allow the user to select different weights for Lévis regression
        weights_levis = [0.40, 0.35, 0.15]
    
        # Perform calculations for Lévis
        rain_data_levis = [rain_today_levis, rain_tomorrow_levis, rain_day_after_levis]
        predicted_flow_levis = calculate_predicted_flow(rain_data_levis, weights_levis)
        st.write(f"Predicted Flow in Lévis: {predicted_flow_levis} m³/s")
    
    # Create a DataFrame with locations, their respective predicted flows, and additional data for visualization
    locations = pd.DataFrame({
        "Location": ["Quebec City", "Levis"],
        "Latitude": [46.808872, 46.693664],
        "Longitude": [-71.316338, -71.070347],
        "Predicted Flow (m³/s)": [predicted_flow_qc, predicted_flow_levis],
        "radius": [predicted_flow_qc * 100, predicted_flow_levis * 100]  # Radius for visualization proportional to flow
    })
    
    # PyDeck layer for flow visualization
    layer = pdk.Layer(
        "ScatterplotLayer",
        locations,
        get_position='[Longitude, Latitude]',
        get_radius='radius',
        get_color=[200, 30, 0, 160],
        pickable=True,
        opacity=0.8,
    )
    
    # Set the view state for the map
    view_state = pdk.ViewState(
        latitude=46.744356,
        longitude=-71.197514,
        zoom=10,
        pitch=0,
    )
    
    # Create and display the PyDeck map
    r = pdk.Deck(layers=[layer], initial_view_state=view_state)
    st.pydeck_chart(r)

    # # Title of the app
    # st.title("🌧️ Survey Planner")
    
    # # Function to perform regression calculation
    # def calculate_predicted_flow(rain_data, weights):
    #     # Assuming a simple linear regression model: flow = sum(rain_data * weights)
    #     return np.sum(np.array(rain_data) * np.array(weights))
    
    # # Sector: Québec
    # st.header("Québec")
    # rain_today_qc = st.slider("Rain Today in Québec (mm)", 0, 100, 25, key="rain_today_qc")
    # rain_tomorrow_qc = st.slider("Predicted Rain Tomorrow in Québec (mm)", 0, 100, 25, key="rain_tomorrow_qc")
    # rain_day_after_qc = st.slider("Predicted Rain 2 Days from Now in Québec (mm)", 0, 100, 25, key="rain_day_after_qc")
    
    # # Regression weights for Québec
    # weights_qc = [0.29949278, 0.29876919, 0.22330731]
    
    # # Perform calculations for Québec
    # rain_data_qc = [rain_today_qc, rain_tomorrow_qc, rain_day_after_qc]
    # predicted_flow_qc = calculate_predicted_flow(rain_data_qc, weights_qc)
    # st.write(f"Predicted Flow in Québec: {predicted_flow_qc} m³/s")
    
    # # Sector: Lévis
    # st.header("Lévis")
    # rain_today_levis = st.slider("Rain Today in Lévis (mm)", 0, 100, 25, key="rain_today_levis")
    # rain_tomorrow_levis = st.slider("Predicted Rain Tomorrow in Lévis (mm)", 0, 100, 25, key="rain_tomorrow_levis")
    # rain_day_after_levis = st.slider("Predicted Rain 2 Days from Now in Lévis (mm)", 0, 100, 25, key="rain_day_after_levis")
    
    # # Allow the user to select different weights for Lévis regression
    # weights_levis = [0.40, 0.35, 0.15]
    
    # # Perform calculations for Lévis
    # rain_data_levis = [rain_today_levis, rain_tomorrow_levis, rain_day_after_levis]
    # predicted_flow_levis = calculate_predicted_flow(rain_data_levis, weights_levis)
    # st.write(f"Predicted Flow in Lévis: {predicted_flow_levis} m³/s")
    
    # # Create a DataFrame with locations, their respective predicted flows, and additional data for visualization
    # locations = pd.DataFrame({
    #     "Location": ["Quebec City", "Levis"],
    #     "Latitude": [46.808872, 46.693664],
    #     "Longitude": [-71.316338, -71.070347],
    #     "Predicted Flow (m³/s)": [predicted_flow_qc, predicted_flow_levis],
    #     "radius": [predicted_flow_qc * 100, predicted_flow_levis * 100]  # Radius for visualization proportional to flow
    # })
    # #46.808872, -71.316338 quebec
    # #46.693664, -71.070347 levis
    # #46.744356, -71.197514 view map
    
    # # PyDeck layer for flow visualization
    # layer = pdk.Layer(
    #     "ScatterplotLayer",
    #     locations,
    #     get_position='[Longitude, Latitude]',
    #     get_radius='radius',
    #     get_color=[200, 30, 0, 160],
    #     pickable=True,
    #     opacity=0.8,
    # )
    
    # # Set the view state for the map
    # view_state = pdk.ViewState(
    #     latitude=46.744356,
    #     longitude=-71.197514,
    #     zoom=10,
    #     pitch=0,
    # )
    
    # # Create and display the PyDeck map
    # r = pdk.Deck(layers=[layer], initial_view_state=view_state)
    # st.pydeck_chart(r)


    
    # st.title("🌧️ Survey Planner")
    # rain_today = st.slider("Rain Today (mm)", 0, 100, 25)
    # rain_tomorrow = st.slider("Predicted Rain Tomorrow (mm)", 0, 100, 25)
    # rain_day_after = st.slider("Predicted Rain 2 Days from Now (mm)", 0, 100, 25)

    # # Placeholder for map and data visualization
    # st.write("Predicted Flow at Different Locations in St. Charles River:")
    
    # # Dummy data for river locations and predicted flow
    # locations = ["Location A", "Location B", "Location C", "Location D"]
    # predicted_flow = np.random.rand(4) * 100  # Random data for demonstration
    
    # # Create a DataFrame for visualization
    # df = pd.DataFrame({
    #     "Location": locations,
    #     "Predicted Flow (m³/s)": predicted_flow
    # })
    
    # # Displaying the data in a table
    # st.dataframe(df)
    
    # # Plotting the data
    # fig, ax = plt.subplots()
    # ax.bar(df["Location"], df["Predicted Flow (m³/s)"])
    # plt.ylabel("Predicted Flow (m³/s)")
    # plt.title("Predicted River Flow at Different Locations")
    # st.pyplot(fig)

    # # Create a DataFrame with locations and their respective predicted flows
    # locations = pd.DataFrame({
    #     "Location": ["Quebec City", "Levis"],
    #     "Latitude": [46.8139, 46.7382],
    #     "Longitude": [-71.2082, -71.2465],
    #     "Predicted Flow (m³/s)": [50, 75]
    # })
    # # Precompute the radius size before defining the PyDeck layer
    # locations['radius'] = locations['Predicted Flow (m³/s)'] * 10
    
    # layer = pdk.Layer(
    #     "ScatterplotLayer",
    #     locations,
    #     get_position='[Longitude, Latitude]',
    #     get_radius='radius',  # Use the precomputed 'radius' column
    #     get_color=[200, 30, 0, 160],
    #     pickable=True,
    #     opacity=0.8,
    # )
    
    # # Set the view state for the map
    # view_state = pdk.ViewState(
    #     latitude=46.8139,  # Centered around the latitude of Quebec City
    #     longitude=-71.2082,  # Centered around the longitude of Quebec City
    #     zoom=8,  # Zoom level adjusted to focus on the area of interest
    #     pitch=0,
    # )
    
    # # Create the PyDeck map
    # r = pdk.Deck(layers=[layer], initial_view_state=view_state)
    
    # # Display the map in the Streamlit app
    # st.pydeck_chart(r)


#Ice analysis"
if choice == "Analyse de la glace - Fr":
    st.title("Analyse des glaces de la station")

    # File upload
    uploaded_file = st.file_uploader("Choose a CSV file")
    if uploaded_file is not None:
        df = pd.read_csv(uploaded_file)
        # Convertir la colonne 'date' en format datetime
        df['date'] = pd.to_datetime(df['date'])
        
        # Définir l'année synthétique et la date synthétique
        df['synthetic_year'] = df['date'].dt.year + (df['date'].dt.month >= 10)
        df['Hiver'] = df['date'].apply(
            lambda x: pd.Timestamp(year=2000 if x.month >= 10 else 2001, 
                                   month=x.month if not (x.month == 2 and x.day == 29) else 3,
                                   day=x.day if not (x.month == 2 and x.day == 29) else 1))
        
        # Filtrer pour n'inclure que d'octobre à mai
        df = df[df['Hiver'].dt.month.isin(list(range(10,13)) + list(range(1,6)))]
        
        # Calculer les degrés-jours de gel
        df['degree_days'] = np.where(df['avg_temperature'] < 0, -df['avg_temperature'], 0)
        
        # Calculer les degrés-jours de gel cumulés par année synthétique
        df = df.sort_values(by=['synthetic_year', 'date'])
        df['cumulative_dd'] = df.groupby('synthetic_year')['degree_days'].cumsum()
        
        # Calculer le nombre de valeurs manquantes par année
        df['missing_values'] = df['avg_temperature'].isna()
        summary = df.groupby('synthetic_year').agg({'cumulative_dd': 'max', 'missing_values': 'sum'})
        
        # Identifier les années synthétiques avec plus de 10 valeurs manquantes
        invalid_years = summary[summary['missing_values'] > 10].index
        
        # Retirer les années avec plus de 10 valeurs manquantes du résumé et du df
        summary = summary[summary['missing_values'] <= 10]
        df = df[~df['synthetic_year'].isin(invalid_years)]
        
        # Créer un DataFrame pour la table
        table_df = summary.reset_index()[['synthetic_year', 'cumulative_dd']]
        
        # Tableau style CSS
        table_df.style.set_properties(**{'background-color': 'lightblue',
                                         'color': 'black',
                                         'border-color': 'white'})
        
        # Convertir les années synthétiques en étiquettes d'hiver
        table_df['synthetic_year'] = [f"{year-1}-{year}" for year in table_df['synthetic_year']]
        
        # Renommer les colonnes pour correspondre aux descriptions fournies
        table_df.columns = ['Hiver', 'Degrés-jours de gel cumulés']
        
        # Afficher la table
        #print(table_df)
        # Display the table
        st.write("Tableau regroupant les degrés-jours cumulés de chaque hiver pour la station")
        st.dataframe(table_df)
        
        ##############################################
        # Histogramme des degrées-jours de gel cumulés
        # Obtenir les étiquettes d'année pour l'histogramme
        year_labels = [f"{year-1}-{year}" for year in summary.index]
        
        fig1, ax = plt.subplots(figsize=(10, 6))
        bars = ax.bar(year_labels, summary['cumulative_dd'], color='steelblue', edgecolor='black')
        
        # Calculer moyenne et écart type
        mean_cddf = summary['cumulative_dd'].mean()
        std_cddf = summary['cumulative_dd'].std()
        st.write("Moyenne")
        st.write(mean_cddf)
        st.write("Écart-type")
        st.write(std_cddf)
        
        # Identifier le nb de données manquantes par hiver
        for bar, missing_values in zip(bars, summary['missing_values']):
            ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height(), str(int(missing_values)), ha='center', va='bottom')
        
        # Ligne de la moyenne
        line_mean = ax.axhline(mean_cddf, color='red', linestyle='--')
        
        # lignes de la déviations standard
        line_std1 = ax.axhline(mean_cddf + std_cddf, color='orange', linestyle='--')
        line_std2 = ax.axhline(mean_cddf - std_cddf, color='orange', linestyle='--')
        
        ax.set_ylabel('Degrés-jours de gel cumulés')
        ax.set_xlabel('Hiver')
        ax.set_title('Histogramme des degrés-jours de gel cumulés par hiver')
        
        # Légende et layout
        legend_labels = ['Moyenne', 'Écart-type +', 'Écart-type -']
        ax.legend([line_mean, line_std1, line_std2], legend_labels, loc='upper right')
        
        plt.xticks(rotation=45)
        plt.tight_layout()
        st.pyplot(fig1)
        
        
        ##################################################
        
        st.write("Le nombre de données quotidiennes manquantes pour chaque période hivernale est identifiées sur les bandes.")
        st.write("  ")
        st.write("  ")
        st.write("Une comparaison de la température moyenne de l'air hivernale par rapport aux degrés-jours de gel cumulés permet d'observé la période moyenne de l'hiver")
        
        # Pour le deuxième graphique
        fig2, ax1 = plt.subplots(figsize=(12, 8))  # Adjust the size as per your requirement
        
        # Dessiner la température moyenne
        df.groupby(df['Hiver'])['avg_temperature'].mean().plot(ax=ax1, color='blue', linewidth=1)
        ax1.set_ylabel('Température moyenne de l\'air (°C)', color='blue')
        ax1.yaxis.grid(True, linestyle='--')
        
        # Dessiner les degrés-jours de gel cumulés
        ax2 = ax1.twinx()
        df.groupby(df['Hiver'])['cumulative_dd'].mean().plot(ax=ax2, color='grey', linewidth=1)
        
        ax2.set_ylabel('Degrés-jours de gel cumulés moyens', color='grey')
        
        # Formater l'axe x pour n'afficher que le premier jour de chaque mois
        ax1.xaxis.set_major_locator(mdates.MonthLocator())
        ax1.xaxis.set_major_formatter(mdates.DateFormatter('%b'))
        plt.setp(ax1.xaxis.get_majorticklabels(), rotation=45)
        
        plt.tight_layout()
        st.pyplot(fig2)
        
        ########################################################################################################
        # Assuming table_df is your DataFrame and 'Degrés-jours de gel cumulés' is the column of interest
        
        data = table_df['Degrés-jours de gel cumulés']
        
        pd.options.display.float_format = '{:.0f}'.format
        
        probabilities = [1/2, 1/5, 1/10, 1/25, 1/50, 1/100]
        recurrence_intervals = [1/prob for prob in probabilities]
        results_df = pd.DataFrame()
        
        distributions_name = ['norm', 'lognorm', 'gumbel_r', 'genextreme']
        
        for dist_name in distributions_name:
            dist = getattr(stats, dist_name)
            params = dist.fit(data)
            
            # Calculate exceedance probabilities
            exceedance_probs = 1 - np.array(probabilities)
            
            quantiles = [dist.ppf(prob, *params[:-2], loc=params[-2], scale=params[-1]) for prob in exceedance_probs]
            row_data = {'Distribution': dist_name}
            for i, interval in enumerate(recurrence_intervals):
                row_data[f'Récurrence {interval}'] = quantiles[i]
            results_df = results_df.append(row_data, ignore_index=True)
        
        results_df.set_index('Distribution', inplace=True)
        
        # Montrer les résultats
        # Arrondir les valeurs à 0 décimal
        results_df = results_df.round(0)
        st.dataframe(results_df)
        
        #Fonction permettant de déterminer les critères AIC et BIC. 
        
        def calculate_aic_bic(data, dist):
            # estimate distribution parameters
            params = dist.fit(data)
        
            # calculate maximum likelihood estimate
            mle = np.sum(dist.logpdf(data, *params))
        
            # calculate number of parameters
            k = len(params)
        
            # calculate AIC and BIC
            aic = 2*k - 2*mle
            bic = np.log(len(data))*k - 2*mle
        
            return aic, bic
        
        # Charger les données
        data = summary['cumulative_dd']
        
        # calculer AIC et BIC pour chaque distribution
        distributions = [norm, lognorm, gumbel_r, genextreme]
        names = ['Normal', 'Lognormal', 'Gumbel', 'GEV']
        
        aic_bic = pd.DataFrame(index=names, columns=['AIC', 'BIC'])
        
        for dist, name in zip(distributions, names):
            aic, bic = calculate_aic_bic(data, dist)
            aic_bic.loc[name, 'AIC'] = aic
            aic_bic.loc[name, 'BIC'] = bic
        
        # Mettre en ordre en fonction du critère BIC car plus critique du nombre de paramètres de la distribution.
        aic_bic = aic_bic.sort_values(by='BIC')
        
        st.dataframe(aic_bic)
        
        
        # identifier 'data' comme étant la valeur CDDF
        data = summary['cumulative_dd'] 
        
        # DataFrame vide pour entreposer les paramètres AIC et BIC. 
        params_df = pd.DataFrame(columns=['Distribution', 'Parameters'])
        
        for dist_name in distributions_name:
            dist = getattr(stats, dist_name)
            params = dist.fit(data)
            
            params_df = params_df.append({'Distribution': dist_name, 'Parameters': params}, ignore_index=True)
        
        st.dataframe(params_df)
        
        fig3, axs = plt.subplots(2, 2, figsize=(12, 12))
        
        axs = axs.ravel()
        
        for ax, dist, name in zip(axs, distributions, names):
            params = dist.fit(data)
            _ = stats.probplot(data, dist=dist, sparams=params, plot=ax)
            ax.set_title(name)
        
        plt.tight_layout()
        st.pyplot(fig3)
        # Dropdown for Stefan's coefficient
        stefans_coefficient = st.selectbox('Stefan\'s Coefficient', [0.7, 0.8, 0.9,1.0,1.1,1.2,1.3,1.4,1.5,1.6,1.7,1.8,1.9,2.0,2.1,2.2,2.3,2.4,2.5,2.6,2.7,2.8,2.9,3.0])
        
        # Dropdown for Effective Resistance of the Ice
        effective_resistance = st.selectbox('Effective Resistance of the Ice', [350, 400, 700, 1100, 1500])
        
        # Dropdown for Slope of the Predicted Riprap
        slope = st.selectbox('Slope of the Predicted Riprap', [0.33, 0.5, 0.66])
        
        # Check if result_df is available
        if 'results_df' in locals():
            # Calculate Theoretical Ice Thickness
            icethickness_df = stefans_coefficient * (results_df ** 0.5)
            st.write("Theoretical Ice Thickness (cm)")
            st.dataframe(icethickness_df)
        
            # Calculate Shear Resistance Dimension
            shear_resistance = (0.0612 * (effective_resistance * slope * (icethickness_df ** 0.5)))/100
            st.write("Shear Resistance Dimension (m)")
            st.dataframe(shear_resistance)

            # Function to create download link for Excel file
            def download_excel_link(excel_file, filename):
                with io.BytesIO() as buffer:
                    excel_file.save(buffer)
                    buffer.seek(0)
                    file = base64.b64encode(buffer.read()).decode('utf-8')
                return f'<a href="data:application/vnd.openxmlformats-officedocument.spreadsheetml.sheet;base64,{file}" download="{filename}">Download Excel file</a>'
            
            # Function to create the Excel workbook
            def create_excel_workbook(df_dict, figure_dict):
                wb = Workbook()
                for sheet_name, df in df_dict.items():
                    ws = wb.create_sheet(title=sheet_name)
                    for r in pd.DataFrame(df).itertuples(index=False, name=None):
                        ws.append(r)
            
                for fig_name, fig in figure_dict.items():
                    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
                    fig.savefig(temp_file.name, format="png")
                    img = Image(temp_file.name)
            
                    ws = wb.create_sheet(title=fig_name)
                    ws.add_image(img, "A1")
            
                    temp_file.close()
            
                return wb
            
            # Assuming df, results_df, icethickness_df, shear_resistance are your DataFrames
            # Assuming fig1, fig2 are your matplotlib figures
            
            # Create a dictionary of DataFrames
            df_dict = {
                "Input Data": df,
                "Results Data": results_df,
                "Ice Thickness": icethickness_df,
                "Shear Resistance": shear_resistance
            }
            
            # Create a dictionary of Figures
            figure_dict = {
                "Figure 1": fig1,
                "Figure 2": fig2
            }
            
            # Streamlit interface
            st.title("Download Excel with Data and Plots")
            if st.button('Generate and Download Excel'):
                wb = create_excel_workbook(df_dict, figure_dict)
                download_link = download_excel_link(wb, "data_and_plots.xlsx")
                st.markdown(download_link, unsafe_allow_html=True)
            
#            def to_excel(df_dict):
#                output = BytesIO()
#                with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
#                    for sheet_name, df in df_dict.items():
#                        df.to_excel(writer, sheet_name=sheet_name, index=False)
#            
#                    # Save figures to a temporary buffer
#                    for fig_name, fig in figure_dict.items():
#                        buffer = BytesIO()
#                        fig.savefig(buffer, format='png')
#                        buffer.seek(0)
#                        image_data = buffer.getvalue()
#                        buffer.close()
#            
#                        # Write the image to a sheet
#                        worksheet = writer.sheets[fig_name]
#                        worksheet.insert_image('A1', fig_name, {'image_data': BytesIO(image_data)})
#            
#                    writer.save()
#                processed_data = output.getvalue()
#                return processed_data
#            
#            # Create a dictionary of DataFrames
#            df_dict = {
#                "Input Data": df,
#                "Results Data": results_df,
#                "Ice Thickness": icethickness_df,
#                "Shear Resistance": shear_resistance
#            }
#            
#            # Create a dictionary of Figures
#            figure_dict = {
#                "Figure 1": fig1,
#                "Figure 2": fig2
#            }

            # Download button
#            if st.button('Download Excel file'):
#                excel_file = to_excel(df_dict)
#                st.download_button(label='📥 Download Excel File',
#                                    data=excel_file,
#                                    file_name='multi_sheet.xlsx',
#                                    mime='application/vnd.ms-excel')

# Ice analysis
if choice == "Ice Analysis - En":
    st.title("❄️ Ice Analysis")

    # File upload
    uploaded_file = st.file_uploader("Choose a CSV file")
    if uploaded_file is not None:
        df = pd.read_csv(uploaded_file)
        # Convert the 'date' column to datetime format
        df['date'] = pd.to_datetime(df['date'])
        
        # Define the synthetic year and synthetic date
        df['synthetic_year'] = df['date'].dt.year + (df['date'].dt.month >= 10)
        df['Winter'] = df['date'].apply(
            lambda x: pd.Timestamp(year=2000 if x.month >= 10 else 2001, 
                                   month=x.month if not (x.month == 2 and x.day == 29) else 3,
                                   day=x.day if not (x.month == 2 and x.day == 29) else 1))
        
        # Filter to include only from October to May
        df = df[df['Winter'].dt.month.isin(list(range(10,13)) + list(range(1,6)))]
        
        # Calculate freeze degree days
        df['degree_days'] = np.where(df['avg_temperature'] < 0, -df['avg_temperature'], 0)
        
        # Calculate cumulative freeze degree days by synthetic year
        df = df.sort_values(by=['synthetic_year', 'date'])
        df['cumulative_dd'] = df.groupby('synthetic_year')['degree_days'].cumsum()
        
        # Calculate the number of missing values per year
        df['missing_values'] = df['avg_temperature'].isna()
        summary = df.groupby('synthetic_year').agg({'cumulative_dd': 'max', 'missing_values': 'sum'})
        
        # Identify synthetic years with more than 10 missing values
        invalid_years = summary[summary['missing_values'] > 10].index
        
        # Remove years with more than 10 missing values from summary and df
        summary = summary[summary['missing_values'] <= 10]
        df = df[~df['synthetic_year'].isin(invalid_years)]
        
        # Create a DataFrame for the table
        table_df = summary.reset_index()[['synthetic_year', 'cumulative_dd']]
        
        # Table CSS style
        table_df.style.set_properties(**{'background-color': 'lightblue',
                                         'color': 'black',
                                         'border-color': 'white'})
        
        # Convert synthetic years to winter labels
        table_df['synthetic_year'] = [f"{year-1}-{year}" for year in table_df['synthetic_year']]
        
        # Rename columns to match provided descriptions
        table_df.columns = ['Winter', 'Cumulative Degree Days of Freezing']
        
        # Display the table
        st.write("Table summarizing the Cumulative Degree Days of Freezing for each winter at the station")
        st.dataframe(table_df)
        
        # Histogram of cumulative freeze degree days
        # Get year labels for the histogram
        year_labels = [f"{year-1}-{year}" for year in summary.index]
        
        fig1, ax = plt.subplots(figsize=(10, 6))
        bars = ax.bar(year_labels, summary['cumulative_dd'], color='steelblue', edgecolor='black')
        
        # Calculate average and standard deviation
        mean_cddf = summary['cumulative_dd'].mean()
        std_cddf = summary['cumulative_dd'].std()
        st.write("Average")
        st.write(mean_cddf)
        st.write("Standard Deviation")
        st.write(std_cddf)
        
        # Indicate the number of missing values per winter
        for bar, missing_values in zip(bars, summary['missing_values']):
            ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height(), str(int(missing_values)), ha='center', va='bottom')
        
        # Average line
        line_mean = ax.axhline(mean_cddf, color='red', linestyle='--')
        
        # Standard deviation lines
        line_std1 = ax.axhline(mean_cddf + std_cddf, color='orange', linestyle='--')
        line_std2 = ax.axhline(mean_cddf - std_cddf, color='orange', linestyle='--')
        
        ax.set_ylabel('Cumulative Degree Days of Freezing')
        ax.set_xlabel('Winter')
        ax.set_title('Histogram of the Cumulative Degree Days of Freezing by Winter')
        
        # Legend and layout
        legend_labels = ['Average', 'Standard Deviation +', 'Standard Deviation -']
        ax.legend([line_mean, line_std1, line_std2], legend_labels, loc='upper right')
        
        plt.xticks(rotation=45)
        plt.tight_layout()
        st.pyplot(fig1)
        
        st.write("The number of daily missing data for each winter period is identified on the bars.")
        st.write(" ")
        st.write(" ")
        st.write("A comparison of the average winter air temperature against the Cumulative Degree Days of Freezing allows observing the average winter period")
        
        # For the second graph
        fig2, ax1 = plt.subplots(figsize=(12, 8))  # Adjust the size as per your requirement
        
        # Plot the average temperature
        df.groupby(df['Winter'])['avg_temperature'].mean().plot(ax=ax1, color='blue', linewidth=1)
        ax1.set_ylabel('Average Air Temperature (°C)', color='blue')
        ax1.yaxis.grid(True, linestyle='--')
        
        # Plot cumulative freeze degree days
        ax2 = ax1.twinx()
        df.groupby(df['Winter'])['cumulative_dd'].mean().plot(ax=ax2, color='grey', linewidth=1)
        
        ax2.set_ylabel('Average Cumulative Degree Days of Freezing', color='grey')
        
        # Format the x-axis to display only the first day of each month
        ax1.xaxis.set_major_locator(mdates.MonthLocator())
        ax1.xaxis.set_major_formatter(mdates.DateFormatter('%b'))
        plt.setp(ax1.xaxis.get_majorticklabels(), rotation=45)
        
        plt.tight_layout()
        st.pyplot(fig2)
        
        ########################################################################################################
        # Assuming table_df is your DataFrame and 'Cumulative Freeze Degree Days' is the column of interest
        
        data = table_df['Cumulative Degree Days of Freezing']
        
        pd.options.display.float_format = '{:.0f}'.format
        
        probabilities = [1/2, 1/5, 1/10, 1/25, 1/50, 1/100]
        recurrence_intervals = [1/prob for prob in probabilities]
        results_df = pd.DataFrame()
        
        distributions_name = ['norm', 'lognorm', 'gumbel_r', 'genextreme']
        
        for dist_name in distributions_name:
            dist = getattr(stats, dist_name)
            params = dist.fit(data)
            
            # Calculate exceedance probabilities
            exceedance_probs = 1 - np.array(probabilities)
            
            quantiles = [dist.ppf(prob, *params[:-2], loc=params[-2], scale=params[-1]) for prob in exceedance_probs]
            row_data = {'Distribution': dist_name}
            for i, interval in enumerate(recurrence_intervals):
                row_data[f'Recurrence {interval}'] = quantiles[i]
            results_df = results_df.append(row_data, ignore_index=True)
        
        results_df.set_index('Distribution', inplace=True)
        
        # Show results
        # Round values to 0 decimal places
        results_df = results_df.round(0)
        st.dataframe(results_df)
        
        # Function to calculate AIC and BIC criteria
        def calculate_aic_bic(data, dist):
            # estimate distribution parameters
            params = dist.fit(data)
        
            # calculate maximum likelihood estimate
            mle = np.sum(dist.logpdf(data, *params))
        
            # calculate number of parameters
            k = len(params)
        
            # calculate AIC and BIC
            aic = 2*k - 2*mle
            bic = np.log(len(data))*k - 2*mle
        
            return aic, bic
        
        # Load the data
        data = summary['cumulative_dd']
        
        # calculate AIC and BIC for each distribution
        distributions = [norm, lognorm, gumbel_r, genextreme]
        names = ['Normal', 'Lognormal', 'Gumbel', 'GEV']
        
        aic_bic = pd.DataFrame(index=names, columns=['AIC', 'BIC'])
        
        for dist, name in zip(distributions, names):
            aic, bic = calculate_aic_bic(data, dist)
            aic_bic.loc[name, 'AIC'] = aic
            aic_bic.loc[name, 'BIC'] = bic
        
        # Order by the BIC criterion as it is more critical of the number of parameters in the distribution.
        aic_bic = aic_bic.sort_values(by='BIC')
        
        st.dataframe(aic_bic)
        
        
        # 'data' identified as the CDDF value
        data = summary['cumulative_dd'] 
        
        # Empty DataFrame to store AIC and BIC parameters
        params_df = pd.DataFrame(columns=['Distribution', 'Parameters'])
        
        for dist_name in distributions_name:
            dist = getattr(stats, dist_name)
            params = dist.fit(data)
            
            params_df = params_df.append({'Distribution': dist_name, 'Parameters': params}, ignore_index=True)
        
        st.dataframe(params_df)
        
        fig3, axs = plt.subplots(2, 2, figsize=(12, 12))
        
        axs = axs.ravel()
        
        for ax, dist, name in zip(axs, distributions, names):
            params = dist.fit(data)
            _ = stats.probplot(data, dist=dist, sparams=params, plot=ax)
            ax.set_title(name)
        
        plt.tight_layout()
        st.pyplot(fig3)
        # Dropdown for Stefan's coefficient
        stefans_coefficient = st.selectbox('Stefan\'s Coefficient', [0.7, 0.8, 0.9, 1.0, 1.1, 1.2, 1.3, 1.4, 1.5, 1.6, 1.7, 1.8, 1.9, 2.0, 2.1, 2.2, 2.3, 2.4, 2.5, 2.6, 2.7, 2.8, 2.9, 3.0])
        
        # Dropdown for Effective Resistance of the Ice
        effective_resistance = st.selectbox('Effective Resistance of the Ice', [350, 400, 700, 1100, 1500])
        
        # Dropdown for Slope of the Predicted Riprap
        slope = st.selectbox('Slope of the Predicted Riprap', [0.33, 0.5, 0.66])
        
        # Check if results_df is available
        if 'results_df' in locals():
            # Calculate Theoretical Ice Thickness
            icethickness_df = stefans_coefficient * (results_df ** 0.5)
            st.write("Theoretical Ice Thickness (cm)")
            st.dataframe(icethickness_df)
        
            # Calculate Shear Resistance Dimension
            shear_resistance = (0.0612 * ((effective_resistance * slope * (icethickness_df/100)) ** 0.5))
            st.write("Shear Resistance Riprap Dimension (m)")
            st.dataframe(shear_resistance)

            # Function to create download link for Excel file
            def download_excel_link(excel_file, filename):
                with io.BytesIO() as buffer:
                    excel_file.save(buffer)
                    buffer.seek(0)
                    file = base64.b64encode(buffer.read()).decode('utf-8')
                return f'<a href="data:application/vnd.openxmlformats-officedocument.spreadsheetml.sheet;base64,{file}" download="{filename}">Download Excel file</a>'
            
            # Function to create the Excel workbook
            def create_excel_workbook(df_dict, figure_dict):
                wb = Workbook()
                for sheet_name, df in df_dict.items():
                    ws = wb.create_sheet(title=sheet_name)
                    for r in pd.DataFrame(df).itertuples(index=False, name=None):
                        ws.append(r)
            
                for fig_name, fig in figure_dict.items():
                    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
                    fig.savefig(temp_file.name, format="png")
                    img = Image(temp_file.name)
            
                    ws = wb.create_sheet(title=fig_name)
                    ws.add_image(img, "A1")
            
                    temp_file.close()
            
                return wb
            
            # Assuming df, results_df, icethickness_df, shear_resistance are your DataFrames
            # Assuming fig1, fig2 are your matplotlib figures
            
            # Create a dictionary of DataFrames
            df_dict = {
                "Input Data": df,
                "Results Data": results_df,
                "Ice Thickness": icethickness_df,
                "Shear Resistance": shear_resistance
            }
            
            # Create a dictionary of Figures
            figure_dict = {
                "Figure 1": fig1,
                "Figure 2": fig2
            }
            
            # Streamlit interface
            st.title("Download Excel with Data and Plots")
            if st.button('Generate and Download Excel'):
                wb = create_excel_workbook(df_dict, figure_dict)
                download_link = download_excel_link(wb, "data_and_plots.xlsx")
                st.markdown(download_link, unsafe_allow_html=True)




#NDBC historical data"
if choice == "NDBC Historical Data Download":
    # Streamlit user interface to input parameters
    st.title("NDBC Historical Data Downloader")
    
    station_id = st.text_input("Enter Station ID", value="44086")
    start_year = st.number_input("Enter Start Year", min_value=1900, max_value=2023, value=2015)
    end_year = st.number_input("Enter End Year", min_value=1900, max_value=2023, value=2020)

    
    df_full_history = download_ndbc_full_history(station_id, start_year, end_year)

    if not df_full_history.empty:
        # Convert date columns to datetime and create a 'Date' column
        df_full_history['Date'] = pd.to_datetime(df_full_history[['YY', 'MM', 'DD']].astype(str).agg('-'.join, axis=1))

        # Select date range for plotting, if 'Date' column is present
        if 'Date' in df_full_history.columns:
            plot_start_date, plot_end_date = st.select_slider(
                "Select Date Range for Plotting",
                options=pd.to_datetime(df_full_history['Date']).sort_values(),
                value=(df_full_history['Date'].min(), df_full_history['Date'].max())
            )
            
            filtered_df = df_full_history[(df_full_history['Date'] >= plot_start_date) & (df_full_history['Date'] <= plot_end_date)]
            
            st.write(filtered_df)
            plot_wave_height(filtered_df, f'Wave Height (WVHT) at Station {station_id} ({plot_start_date.date()} to {plot_end_date.date()})')
        else:
            st.write("Date column not found in the DataFrame.")
    else:
        st.write("No data available for the specified range.")
    
    # if st.button("Download Data"):
    #     df_full_history = download_ndbc_full_history(station_id, start_year, end_year)
    
    #     # Display the DataFrame in the Streamlit app
    #     if not df_full_history.empty:
    #         st.write(df_full_history)
    #         # Select date range for plotting
    #         plot_start_date, plot_end_date = st.select_slider(
    #             "Select Date Range for Plotting",
    #             options=pd.to_datetime(df_full_history['Date']).sort_values(),
    #             value=(df_full_history['Date'].min(), df_full_history['Date'].max())
    #         )
            
    #         filtered_df = df_full_history[(df_full_history['Date'] >= plot_start_date) & (df_full_history['Date'] <= plot_end_date)]
            
    #         st.write(filtered_df)
    #         plot_wave_height(filtered_df, f'Wave Height (WVHT) at Station {station_id} ({plot_start_date.date()} to {plot_end_date.date()})')
    #         #plot_wave_height(df_full_history, station_id)
    #     else:
    #         st.write("No data available for the specified range.")
            
#"Water level CEHQ"
if choice == "Water level CEHQ":
    # Streamlit app layout
    st.title("Hydrometric Station Data Processor")
    
    # User input for station number
    station_number = st.text_input("Enter the Station Number", "030240_N")
    
    # URL construction based on station number
    url = f"https://www.cehq.gouv.qc.ca/depot/historique_donnees/fichier/{station_number}.txt"
    
     # Function to fetch and process data
    def fetch_and_process_data(url):
        response = requests.get(url)
        lines = response.text.split('\n')
    
        # Extracting station description
        description = lines[:21]
        
        # Processing the data
        data = [line.split() for line in lines[22:] if line.strip()]  # Assuming data starts from line 22
        df = pd.DataFrame(data)

        df.columns = ['Station', 'Date', 'Water Level', 'Info']

        # Splitting 'Column2' into 'year', 'month', and 'day'
        df[['year', 'month', 'day']] = df['Date'].str.split('/', expand=True)
        df['Water Level'] = pd.to_numeric(df['Water Level'], errors='coerce')

        
        # Convert year, month, and day to integers
        df['year'] = df['year'].astype(int)
        df['month'] = df['month'].astype(int)
        df['day'] = df['day'].astype(int)

        # Group by year and calculate statistics
        grouped = df.groupby('year')['Water Level']
        max_indices = grouped.idxmax()
        min_indices = grouped.idxmin()
        max_values = grouped.max()
        min_values = grouped.min()
        none_counts = grouped.apply(lambda x: x.isna().sum())

        # Extracting dates for max and min values
        #max_dates = df.loc[max_indices, 'Date']
        #min_dates = df.loc[min_indices, 'Date']
        
        # Combining results
        annual_stats = pd.DataFrame({
            'None Count': none_counts,
            'Max Value':grouped.max(),
            #'Date of Max Value': df.loc[grouped.idxmax(), 'Date'],
            'Min Value': grouped.min(),
            #'Date of Min Value':  df.loc[grouped.idxmin(), 'Date']
        })
        date_stats = pd.DataFrame({
            'Date of Max Value': df.loc[grouped.idxmax(), 'Date'],
            'Date of Min Value':  df.loc[grouped.idxmin(), 'Date']
         })
        
        # # Calculating annual statistics
        # none_counts = df.groupby('year')['Water Level'].apply(lambda x: x.isna().sum())  
        # max_values = df.groupby('year')['Water Level'].max()
        # min_values = df.groupby('year')['Water Level'].min()
        
        # annual_stats = pd.DataFrame({
        #     'None Count': none_counts,
        #     'Max Value': max_values,
        #     'Min Value': min_values
        # })

        
        # if len(df.columns) == 4:
        #     df.columns = ['Column1', 'Column2', 'Column3', 'Column4']  # Replace with actual column names
        # else:
        #     st.error(f"Unexpected number of columns. Found: {len(df.columns)}")
        #     return None, None, None
    
        # # Handling the date column
        # df['Date'] = pd.to_datetime(df['Date'], errors='coerce')  # Convert Date column to datetime
        
        ## Convert data columns to numeric as needed
        #df['Column2'] = pd.to_numeric(df['Column2'], errors='coerce')  # Example for numeric conversion
    
        # Calculating annual min, max, and missing values
        #annual_stats = df.agg({'Column2': ['min', 'max'], 'Column3': ['count']})
        
        return description, df, annual_stats,date_stats
        
        #return description, df
    
    # Display the results
    if st.button("Fetch Data"):
        #description, df = fetch_and_process_data(url)
        description, df, annual_stats, date_stats = fetch_and_process_data(url)
        st.write("Station Description:")
        st.write(description)
        st.write("Dataframe:")
        st.dataframe(df)
        st.write("Annual Statistics:")
        st.dataframe(annual_stats)

        # Set the style for nicer plots
        plt.style.use('seaborn-darkgrid')
        
        # Creating a figure and axis
        fig, ax = plt.subplots(figsize=(10, 6))
        
        # Width of a bar 
        width = 0.25       
        
        # Setting the positions of the bars
        ind = annual_stats.index.astype(str)  # Assuming the index of annual_stats is the year
        ind = range(len(ind))  # Convert to numeric index for plotting
        
        # Plotting
        ax.bar(ind, annual_stats['Max Value'], width, label='Max Value')
        ax.bar([i + width for i in ind], annual_stats['Min Value'], width, label='Min Value')
        ax.bar([i + width*2 for i in ind], annual_stats['None Count'], width, label='None Count')
        
        # Adding labels and title
        ax.set_xlabel('Year')
        ax.set_ylabel('Values')
        ax.set_title('Annual Water Level Statistics')
        ax.set_xticks([i + width for i in ind])
        ax.set_xticklabels(annual_stats.index.astype(str))
        
        # Adding a legend
        ax.legend()
        
        # Show the plot
        st.pyplot(fig)

        st.dataframe(date_stats)
    
        # # Plotting
        # fig, ax = plt.subplots()
        # ax.plot(df['Column1'], df['Column2'])  # Replace 'Column1' and 'Column2' with actual columns
        # ax.set_title("Annual Min and Max")
        # st.pyplot(fig)
    
# # Home page
# if choice == "Home":
#     st.title("Water Engineering Tools")
#     st.write(
#         """
#         Welcome to the Water Engineering Tools web app created by a junior engineer.
#         This web app includes the following tools:

#         1. Hydrograph Producer: This tool allows you to import a CSV file containing daily flow data time series and plots the hydrograph for each year. It also provides the maximum, minimum, and number of missing values.

#         2. Peak Flow Comparison: This tool compares two time series. The first time series contains the daily flow data of a river, while the second contains flow data for every 15 minutes of the same river. The tool compares the maximum value for each year of both time series and returns a table with all the ratios for each specific year. The last row displays the mean of these ratios.

#         3. Camera Viewer: This tool allows you to input images and displays the image on the webpage.

#         4. Frequency Analysis: This tool performs frequency analysis on the maximum flow data using various probability distributions and generates a Word document with the analysis results.
#         """
#     )
if choice == "Home":
    st.title("💧 Water Engineering Tools 🌊")
    st.write(
        """
        Welcome to the 💡 Water Engineering Tools web app created by a junior engineer.
        Dive into a world of water engineering with the following tools:

        **Available Tools:**

        - 🌊 **Hydrograph Producer**: 
        This tool allows you to import a CSV file containing daily flow data time series and plots the hydrograph for each year. It also provides the maximum, minimum, and number of missing values.
          * Import CSV with daily flow data.
          * Plot hydrographs for each year.
          * View max, min, and missing values.

        - ❄️ **Ice Analysis** (English Version): 
        Based on meteorological data, this tool determines the theoretical ice thickness and riprap dimension. It offers valuable insights into ice behavior under various environmental conditions.
          * Analyze meteorological data.
          * Determine ice thickness and riprap dimension.
          * Understand ice behavior under different conditions.

        - 🌧️ **Survey Planner**: 
        Utilizing predicted rain events in a predictive model, this tool forecasts the flow in specific rivers. It's an essential tool for managers to plan surveys, ensuring effective resource allocation for survey teams.
          * Predict river flow using rain event models.
          * Plan surveys efficiently.
          * Essential for resource allocation and team management.

        - 🥨 **CrissPy**: 
        CrissPy stands out as a tool designed to process and analyze hydraulic node data effectively from a CRISSP model using Python. It caters to professionals in water engineering and environmental sciences, offering insightful time series analysis.
          * Upload a zip file containing multiple .hdw files.
          * Allows users to specify and focus on a particular node for detailed analysis.
          * Enable users to focus more on analysis and less on data processing complexities.
        """
    )


if choice == "Hydrograph Producer":
    st.header("🌊 Hydrograph Producer")

    uploaded_file = st.file_uploader(
        "Upload a CSV file with daily flow data (Date, Flow, Year):", type="csv")
    if uploaded_file is not None:
        daily_flow_data = pd.read_csv(uploaded_file)
        daily_flow_data['Date'] = pd.to_datetime(daily_flow_data['Date'])
        daily_flow_data.set_index('Date', inplace=True)

        sep_day = st.number_input(
            "Separation Day (default: 1):", min_value=1, max_value=31, value=1)
        sep_month = st.number_input(
            "Separation Month (default: 7):", min_value=1, max_value=12, value=7)
        spring_volume_period = st.number_input(
            "Spring Volume Period (default: 30):", min_value=1, max_value=365, value=30)
        fall_volume_period = st.number_input(
            "Fall Volume Period (default: 10):", min_value=1, max_value=365, value=10)

        if st.button("Generate Hydrographs and Tables"):
            wb = generate_hydrographs_and_tables(
                daily_flow_data, sep_day, sep_month, spring_volume_period, fall_volume_period)
            st.markdown(download_link(
                wb, "hydrograph_analysis.xlsx"), unsafe_allow_html=True)

        # Convert the "Date" column to a datetime object
        # df = daily_flow_data
        # df["Date"] = pd.to_datetime(df["Date"])
        # years = df["Year"].unique()
        st.subheader("Hydrographs")
        years = daily_flow_data["Year"].unique()
        for year in years:
            df_year = daily_flow_data[daily_flow_data["Year"] == year]
            #df_year.set_index("Date", inplace=True)
            # Calculate the rolling sum of flow for spring and fall periods
            df_year["Rolling_Spring"] = df_year.loc[:pd.Timestamp(
                year, sep_month, sep_day), "Flow"].rolling(window=spring_volume_period).sum()
            df_year["Rolling_Fall"] = df_year.loc[pd.Timestamp(
                year, sep_month, sep_day):, "Flow"].rolling(window=fall_volume_period).sum()

            # Find the maximum rolling sum periods for spring and fall
            spring_start_date = df_year["Rolling_Spring"].idxmax(
            ) - pd.Timedelta(days=spring_volume_period - 1)
            spring_end_date = df_year["Rolling_Spring"].idxmax()
            fall_start_date = df_year["Rolling_Fall"].idxmax(
            ) - pd.Timedelta(days=fall_volume_period - 1)
            fall_end_date = df_year["Rolling_Fall"].idxmax()

            fig, ax = plt.subplots(figsize=(15, 6))
            ax.plot(df_year.index, df_year["Flow"])
            ax.axvline(pd.Timestamp(year, sep_month, sep_day),
                       color="black", linestyle="--", label="Separation Date")

            # Highlight the spring and fall volume periods in red and green, respectively
            ax.axvspan(spring_start_date, spring_end_date, alpha=0.3,
                       color="red", label="Spring Volume Period")
            ax.axvspan(fall_start_date, fall_end_date, alpha=0.3,
                       color="green", label="Fall Volume Period")

            # Maximum and minimum values for spring and fall periods
            spring_max = df_year.loc[:pd.Timestamp(
                year, sep_month, sep_day), "Flow"].max()
            spring_min = df_year.loc[:pd.Timestamp(
                year, sep_month, sep_day), "Flow"].min()
            fall_max = df_year.loc[pd.Timestamp(
                year, sep_month, sep_day):, "Flow"].max()
            fall_min = df_year.loc[pd.Timestamp(
                year, sep_month, sep_day):, "Flow"].min()

            # Add red and green dots for maximum and minimum values of spring and fall periods
            spring_max_date = df_year.loc[:pd.Timestamp(
                year, sep_month, sep_day), "Flow"].idxmax()
            spring_min_date = df_year.loc[:pd.Timestamp(
                year, sep_month, sep_day), "Flow"].idxmin()
            fall_max_date = df_year.loc[pd.Timestamp(
                year, sep_month, sep_day):, "Flow"].idxmax()
            fall_min_date = df_year.loc[pd.Timestamp(
                year, sep_month, sep_day):, "Flow"].idxmin()

            ax.plot(spring_max_date, spring_max, "ro")
            ax.plot(spring_min_date, spring_min, "go")
            ax.plot(fall_max_date, fall_max, "ro")
            ax.plot(fall_min_date, fall_min, "go")

            ax.set_title(f"Hydrograph for {year}")
            ax.set_ylabel("Flow")
            ax.legend(loc="best")
            ax.xaxis.set_major_formatter(mdates.DateFormatter("%b"))

            st.pyplot(fig)

            st.write(
                f"Spring Volume Period: {spring_start_date.strftime('%d-%m')} to {spring_end_date.strftime('%d-%m')}")
            st.write(
                f"Fall Volume Period: {fall_start_date.strftime('%d-%m')} to {fall_end_date.strftime('%d-%m')}")
            st.write(
                f"Max and Min for Spring: {spring_max} ({spring_max_date.strftime('%d-%m')}), {spring_min} ({spring_min_date.strftime('%d-%m')})")
            st.write(
                f"Max and Min for Fall: {fall_max} ({fall_max_date.strftime('%d-%m')}), {fall_min} ({fall_min_date.strftime('%d-%m')})")
    else:
        st.info("Please upload a CSV file.")

# Peak Flow Comparison page
elif choice == "Peak Flow Comparison":
    st.title("Peak Flow Comparison")
    st.write("Make sure these csv files contain matching years.")

    uploaded_file1 = st.file_uploader(
        "Choose the first CSV file (daily flow data)", type="csv")
    uploaded_file2 = st.file_uploader(
        "Choose the second CSV file (flow data every 15 minutes)", type="csv")

    if uploaded_file1 is not None and uploaded_file2 is not None:
        df1 = pd.read_csv(uploaded_file1)
        df2 = pd.read_csv(uploaded_file2)
    
        max_df1 = df1.groupby("Year")["Flow"].max().reset_index().rename(columns={"Flow": "max_Daily"})
        max_df2 = df2.groupby("Year")["Flow"].max().reset_index().rename(columns={"Flow": "max_Instant"})
    
        merged_df = pd.merge(max_df1, max_df2, on="Year", how="inner")
        merged_df["Ratio"] = merged_df["max_Instant"] / merged_df["max_Daily"]
        mean_ratio = merged_df["Ratio"].mean()
    
        st.write("Details for each year:")
        st.write(merged_df)
    
        st.write(f"Mean of ratios: {mean_ratio}")

# # EC Canada Daily Data
elif choice == "EC Daily Data Analysis":
    # Set the style for the plots
    plt.style.use('seaborn-whitegrid')
    
    st.title('Climate Data Analysis')

    # File uploader
    uploaded_file = st.file_uploader("Upload a CSV file with climate data:", type="csv")
    
    # Only proceed with the rest of the app if a file is uploaded
    if uploaded_file is not None:
        df = process_file(uploaded_file)
        if df is not None:
             # Text input for start date
            start_date = st.text_input('Enter start date (YYYY-MM-DD):')
            # Text input for end date
            end_date = st.text_input('Enter end date (YYYY-MM-DD):')
    
            # Check if the dates are entered and valid
            if start_date and end_date:
                try:
                    start_date = pd.to_datetime(start_date)
                    end_date = pd.to_datetime(end_date)
                    
                    # Filter the DataFrame based on the input dates
                    filtered_df = df[(df['Date'] >= start_date) & (df['Date'] <= end_date)]

                    st.subheader('Temperature Analysis')
                    fig, ax = plt.subplots(figsize=(10, 5))
                    
                    # Remove grid and add black contour
                    ax.grid(True, linestyle='--')
                    ax.spines['top'].set_color('black')
                    ax.spines['bottom'].set_color('black')
                    ax.spines['left'].set_color('black')
                    ax.spines['right'].set_color('black')

                    # Setting the ticks
                    ax.tick_params(axis='x', which='both', bottom=True, top=True, labelbottom=True)
                    ax.tick_params(axis='y', which='both', left=True, right=True, labelleft=True)

                    # Plotting the line plots for temperature
                    ax.plot(filtered_df['Date'], filtered_df['Max Temp (°C)'], label='Max Temp', color='darkred')
                    ax.plot(filtered_df['Date'], filtered_df['Min Temp (°C)'], label='Min Temp', color='red')
                    ax.plot(filtered_df['Date'], filtered_df['Mean Temp (°C)'], label='Mean Temp', color='salmon')
                    
                    # Calculate the monthly average temperatures
                    monthly_avg_temps = filtered_df.resample('M', on='Date')['Mean Temp (°C)'].mean()
                    month_names = monthly_avg_temps.index.strftime('%B')

                    # Creating a table at the top of the graph
                    col_labels = ['Month', 'Avg Temp (°C)']
                    table_vals = list(zip(month_names, monthly_avg_temps.round(1)))
                    table = ax.table(cellText=table_vals, colLabels=col_labels, loc='top', cellLoc='center')
                    table.auto_set_font_size(False)
                    table.set_fontsize(9)
                    table.scale(1, 1.5)

                    # Overlay a bar chart for monthly average temperatures
                    ax.bar(monthly_avg_temps.index, monthly_avg_temps, width=20, color='pink', label='Monthly Avg Temp', alpha=0.5, align='center')
                    
                    #ax.set_xlabel('Date')
                    ax.set_ylabel('Temperature (°C)')
                    ax.xaxis.set_major_formatter(mdates.DateFormatter('%Y-%m-%d'))
                    ax.xaxis.set_major_locator(mdates.AutoDateLocator())  # Automatic tick locator
                    fig.autofmt_xdate()

                    # Adjusting the legend
                    ax.legend(loc='upper left', bbox_to_anchor=(1, 1), fontsize='small', frameon=True)

                    # Adjust layout to make room for the table and legend
                    plt.subplots_adjust(left=0.2, bottom=0.2, top=0.8, right=0.8)
                    st.pyplot(fig)

                    
                    # fig, ax = plt.subplots(figsize=(10, 5))
                    
                    # # Calculate the monthly average temperatures
                    # monthly_avg_temps = filtered_df.resample('M', on='Date')['Mean Temp (°C)'].mean()
                    # month_names = monthly_avg_temps.index.strftime('%B')

                    # # Creating a table at the top of the graph
                    # col_labels = ['Month', 'Avg Temp (°C)']
                    # table_vals = list(zip(month_names, monthly_avg_temps.round(1)))
                    # table = ax.table(cellText=table_vals, colLabels=col_labels, loc='top', cellLoc='center')
                    # table.auto_set_font_size(False)
                    # table.set_fontsize(9)
                    # table.scale(1, 1.5)

                    # # Plotting the line plots for temperature
                    # ax.plot(filtered_df['Date'], filtered_df['Max Temp (°C)'], label='Max Temp', color='darkred')
                    # ax.plot(filtered_df['Date'], filtered_df['Min Temp (°C)'], label='Min Temp', color='red')
                    # ax.plot(filtered_df['Date'], filtered_df['Mean Temp (°C)'], label='Mean Temp', color='salmon')
                    
                    # # Overlay a bar chart for monthly average temperatures
                    # ax.bar(monthly_avg_temps.index, monthly_avg_temps, width=20, color='pink', label='Monthly Avg Temp', alpha=0.5, align='center')
                    
                    # ax.set_xlabel('Date')
                    # ax.set_ylabel('Temperature (°C)')
                    # ax.xaxis.set_major_formatter(mdates.DateFormatter('%Y-%m-%d'))
                    # ax.xaxis.set_major_locator(mdates.AutoDateLocator())  # Automatic tick locator
                    # fig.autofmt_xdate()
                    # ax.legend()

                    # # Adjust layout to make room for the table
                    # plt.subplots_adjust(left=0.2, bottom=0.2, top=0.8)
                    # st.pyplot(fig)

    
                    # st.subheader('Temperature Analysis')
                    # fig, ax = plt.subplots(figsize=(10, 5))
                    
                    # # Plotting the line plots for temperature
                    # ax.plot(filtered_df['Date'], filtered_df['Max Temp (°C)'], label='Max Temp', color='darkred')
                    # ax.plot(filtered_df['Date'], filtered_df['Min Temp (°C)'], label='Min Temp', color='red')
                    # ax.plot(filtered_df['Date'], filtered_df['Mean Temp (°C)'], label='Mean Temp', color='salmon')
                    
                    # # Calculate the monthly average temperatures
                    # monthly_avg_temps = filtered_df.resample('M', on='Date')['Mean Temp (°C)'].mean()
                    
                    # # Overlay a bar chart for monthly average temperatures
                    # ax.bar(monthly_avg_temps.index, monthly_avg_temps, width=20, color='pink', label='Monthly Avg Temp', alpha=0.5, align='center')
                    
                    # # Annotate the bars with values
                    # for idx, value in enumerate(monthly_avg_temps):
                    #     ax.annotate(f'{value:.1f}', 
                    #                 (monthly_avg_temps.index[idx], value), 
                    #                 textcoords="offset points", 
                    #                 xytext=(0,10), 
                    #                 ha='center')
                    
                    # ax.set_xlabel('Date')
                    # ax.set_ylabel('Temperature (°C)')
                    # ax.xaxis.set_major_formatter(mdates.DateFormatter('%Y-%m-%d'))
                    # ax.xaxis.set_major_locator(mdates.AutoDateLocator())  # Automatic tick locator
                    # fig.autofmt_xdate()
                    # ax.legend()
                    # st.pyplot(fig)
                    
                    # fig, ax = plt.subplots(figsize=(8, 4))
                    # ax.plot(filtered_df['Date'], filtered_df['Max Temp (°C)'], label='Max Temp', color='tomato')
                    # ax.plot(filtered_df['Date'], filtered_df['Min Temp (°C)'], label='Min Temp', color='dodgerblue')
                    # ax.plot(filtered_df['Date'], filtered_df['Mean Temp (°C)'], label='Mean Temp', color='green')
                    # ax.set_xlabel('Date')
                    # ax.set_ylabel('Temperature (°C)')
                    # ax.xaxis.set_major_formatter(mdates.DateFormatter('%Y-%m-%d'))
                    # fig.autofmt_xdate()
                    # ax.legend()
                    # st.pyplot(fig)

                    # Precipitation Analysis
                    st.subheader('Precipitation Analysis')
                    fig, ax = plt.subplots(figsize=(10, 5))

                    # Remove grid and add black contour
                    ax.grid(True, linestyle='--')
                    ax.spines['top'].set_color('black')
                    ax.spines['bottom'].set_color('black')
                    ax.spines['left'].set_color('black')
                    ax.spines['right'].set_color('black')


                    # Calculate the monthly precipitation stats
                    monthly_precip = filtered_df.resample('M', on='Date')['Total Precip (mm)']
                    monthly_stats = pd.DataFrame({
                        'Min Rain (mm)': monthly_precip.min(),
                        'Max Rain (mm)': monthly_precip.max(),
                        'Total Rain (mm)': monthly_precip.sum()
                    }).round(1)
                    month_names = monthly_stats.index.strftime('%B')

                    # Creating a table at the top of the graph
                    table_vals = [month_names] + [monthly_stats[col].values for col in monthly_stats.columns]
                    table = ax.table(cellText=table_vals, rowLabels=['Month', 'Min Rain', 'Max Rain', 'Total Rain'], loc='top', cellLoc='center')
                    table.auto_set_font_size(False)
                    table.set_fontsize(6)  # Adjusted font size
                    table.scale(1, 1.5)

                    # Plotting the bar chart for precipitation
                    ax.bar(filtered_df['Date'], filtered_df['Total Precip (mm)'], color='deepskyblue', width=1.0)  # Adjusted color and width
                    ax.set_xlabel('Date')
                    ax.set_ylabel('Precipitation (mm)')
                    ax.xaxis.set_major_formatter(mdates.DateFormatter('%Y-%m-%d'))
                    ax.xaxis.set_major_locator(mdates.AutoDateLocator())  # Automatic tick locator
                    fig.autofmt_xdate()

                    # Adjust layout to make room for the table
                    plt.subplots_adjust(left=0.2, bottom=0.2, top=0.8)
                    st.pyplot(fig)

                    
                    # fig, ax = plt.subplots(figsize=(10, 5))

                    # # Calculate the monthly precipitation stats
                    # monthly_precip = filtered_df.resample('M', on='Date')['Total Precip (mm)']
                    # monthly_stats = pd.DataFrame({
                    #     'Min Rain (mm)': monthly_precip.min(),
                    #     'Max Rain (mm)': monthly_precip.max(),
                    #     'Total Rain (mm)': monthly_precip.sum()
                    # }).round(1)
                    # month_names = monthly_stats.index.strftime('%B')

                    # # Creating a table at the top of the graph
                    # table_vals = [month_names] + [monthly_stats[col].values for col in monthly_stats.columns]
                    # table = ax.table(cellText=table_vals, rowLabels=['Month', 'Min Rain', 'Max Rain', 'Total Rain'], loc='top', cellLoc='center')
                    # table.auto_set_font_size(False)
                    # table.set_fontsize(9)
                    # table.scale(1, 1.5)

                    # # Plotting the bar chart for precipitation
                    # ax.bar(filtered_df['Date'], filtered_df['Total Precip (mm)'], color='lightblue')
                    # ax.set_xlabel('Date')
                    # ax.set_ylabel('Precipitation (mm)')
                    # ax.xaxis.set_major_formatter(mdates.DateFormatter('%Y-%m-%d'))
                    # ax.xaxis.set_major_locator(mdates.AutoDateLocator())  # Automatic tick locator
                    # fig.autofmt_xdate()

                    # # Adjust layout to make room for the table
                    # plt.subplots_adjust(left=0.2, bottom=0.2, top=0.8)
                    # st.pyplot(fig)
                    
                    # fig, ax = plt.subplots(figsize=(8, 4))
                    # ax.bar(filtered_df['Date'], filtered_df['Total Precip (mm)'], color='lightblue')
                    # ax.set_xlabel('Date')
                    # ax.set_ylabel('Precipitation (mm)')
                    # ax.xaxis.set_major_formatter(mdates.DateFormatter('%Y-%m-%d'))
                    # fig.autofmt_xdate()
                    # st.pyplot(fig)

                    # Snow Analysis
                    st.subheader('Snow Analysis')
                    fig, ax = plt.subplots(figsize=(8, 4))
                    ax.plot(filtered_df['Date'], filtered_df['Snow on Grnd (cm)'], color='lightgrey')
                    ax.set_xlabel('Date')
                    ax.set_ylabel('Snow on Ground (cm)')
                    ax.xaxis.set_major_formatter(mdates.DateFormatter('%Y-%m-%d'))
                    fig.autofmt_xdate()
                    st.pyplot(fig)

                    ##For now, not added
                    # # Wind Gust Analysis
                    # st.subheader('Wind Gust Analysis')
                    # fig, ax = plt.subplots(figsize=(8, 4))
                    # ax.plot(filtered_df['Date'], filtered_df['Spd of Max Gust (km/h)'], color='purple')
                    # ax.set_xlabel('Date')
                    # ax.set_ylabel('Speed of Max Gust (km/h)')
                    # ax.xaxis.set_major_formatter(mdates.DateFormatter('%Y-%m-%d'))
                    # fig.autofmt_xdate()
                    # st.pyplot(fig)

                    
                    # Temperature analysis
                    # st.subheader('Temperature Analysis')
                    # fig, ax = plt.subplots()
                    # ax.plot(filtered_df['Date/Time'], filtered_df['Max Temp (°C)'], label='Max Temp')
                    # ax.plot(filtered_df['Date/Time'], filtered_df['Min Temp (°C)'], label='Min Temp')
                    # ax.plot(filtered_df['Date/Time'], filtered_df['Mean Temp (°C)'], label='Mean Temp')
                    # ax.set_xlabel('Date')
                    # ax.set_ylabel('Temperature (°C)')
                    # ax.legend()
                    # st.pyplot(fig)
                
                    # Precipitation analysis
                    # st.subheader('Precipitation Analysis')
                    # fig, ax = plt.subplots()
                    # ax.bar(filtered_df['Date/Time'], filtered_df['Total Precip (mm)'])
                    # ax.set_xlabel('Date')
                    # ax.set_ylabel('Precipitation (mm)')
                    # st.pyplot(fig)
                
                    # Snow analysis
                    # st.subheader('Snow Analysis')
                    # fig, ax = plt.subplots()
                    # ax.plot(filtered_df['Date/Time'], filtered_df['Snow on Grnd (cm)'])
                    # ax.set_xlabel('Date')
                    # ax.set_ylabel('Snow on Ground (cm)')
                    # st.pyplot(fig)
                
                    # Wind Gust analysis
                    # st.subheader('Wind Gust Analysis')
                    # fig, ax = plt.subplots()
                    # ax.plot(filtered_df['Date/Time'], filtered_df['Spd of Max Gust (km/h)'])
                    # ax.set_xlabel('Date')
                    # ax.set_ylabel('Speed of Max Gust (km/h)')
                    # st.pyplot(fig)
                except ValueError as e:
                        st.error("The dates entered are invalid. Please enter valid dates in the format YYYY-MM-DD.")
    
# # Camera Viewer page

elif choice == "Camera Viewer":
    st.title("Camera Viewer")

    # Step 1: Add a file uploader for the zip file
    uploaded_file = st.file_uploader("Choose a zip file", type="zip")

    if uploaded_file is not None:
        # Step 2: Extract the zip file and read its contents
        with zipfile.ZipFile(uploaded_file, "r") as zfile:
            zfile.extractall("temp_folder")

        # Step 3: Process the image files and the CSV files
        image_files = []
        hydrograph_df = None
        rain_df = None
        temperature_df = None

        for root, _, files in os.walk("temp_folder"):
            for file in files:
                if file.lower().endswith(".jpg"):
                    image_files.append(os.path.join(root, file))
                elif file == "Hydrograph.csv":
                    hydrograph_df = pd.read_csv(os.path.join(
                        root, file), parse_dates=[["Date", "Time"]])
                elif file == "Rain.csv":
                    rain_df = pd.read_csv(os.path.join(
                        root, file), parse_dates=[["Date", "Time"]])
                elif file == "Temperature.csv":
                    temperature_df = pd.read_csv(os.path.join(
                        root, file), parse_dates=[["Date", "Time"]])

        st.write(f"Uploaded file: {uploaded_file}")
        st.write(f"Number of image files: {len(image_files)}")
        st.write(f"Hydrograph dataframe: {hydrograph_df.head()}")
        st.write(f"Rain dataframe: {rain_df.head()}")
        st.write(f"Temperature dataframe: {temperature_df.head()}")

        # Step 4: For each image file, display the image and plot the graphs
        for img_file in image_files:
            # Extract the time from the image filename
            # Remove ".jpg" from the filename
            img_time_str = os.path.basename(img_file)[:-4]
            img_time = datetime.strptime(img_time_str, "%m%d%Y%H%M")

            st.image(
                img_file, caption=f"Image taken at {img_time.strftime('%Y-%m-%d %H:%M')}")

            fig, ax = plt.subplots(3, 1, figsize=(10, 15), sharex=True)

            for idx, (df, title, ylabel) in enumerate(zip([hydrograph_df, rain_df, temperature_df], ["Flow", "Rain", "Temperature"], ["Flow", "Rain", "Temperature"])):
                # Filter the data within 3 days before and after the image time
                mask = (df["Date_Time"] >= img_time - timedelta(days=3)
                        ) & (df["Date_Time"] <= img_time + timedelta(days=3))
                data = df[mask]

                # Plot the graph
                ax[idx].plot(data["Date_Time"], data[title], label=title)
                ax[idx].set_ylabel(ylabel)

                if not data.empty:  # Add this condition before finding the closest time
                    # Add a red dot at the image time
                    min_index = (data["Date_Time"] - img_time).abs().idxmin()
                    if min_index in data.index:
                        closest_time = data.loc[min_index, "Date_Time"]
                    else:
                        closest_time = img_time
                    ax[idx].plot(closest_time, data.loc[data["Date_Time"] ==
                                                        closest_time, title].values[0], "ro", label="Image time")

                    # Adjust y-axis label to 6 increments
                    min_value, max_value = ax[idx].get_ylim()
                    ax[idx].yaxis.set_ticks(
                        np.linspace(min_value, max_value, 6))

                ax[idx].legend()
            # for idx, (df, title, ylabel) in enumerate(zip([hydrograph_df, rain_df, temperature_df], ["Flow", "Rain", "Temperature"], ["Flow", "Rain", "Temperature"])):
            #     # Filter the data within 15 days before and after the image time
            #     mask = (df["Date_Time"] >= img_time - timedelta(days=15)
            #             ) & (df["Date_Time"] <= img_time + timedelta(days=15))
            #     data = df[mask]

            #     # Plot the graph
            #     ax[idx].plot(data["Date_Time"], data[title], label=title)
            #     ax[idx].set_ylabel(ylabel)

            #     if not data.empty:  # Add this condition before finding the closest time
            #         # Add a red dot at the image time
            #         min_index = (data["Date_Time"] - img_time).abs().idxmin()
            #         if min_index in data.index:
            #             closest_time = data.loc[min_index, "Date_Time"]
            #         else:
            #             closest_time = img_time
            #         ax[idx].plot(closest_time, data.loc[data["Date_Time"] ==
            #                                             closest_time, title].values[0], "ro", label="Image time")

            #     ax[idx].legend()

            plt.xlabel("Date")
            plt.xticks(rotation=45)
            st.pyplot(fig)
        shutil.rmtree("temp_folder")
# elif choice == "Camera Viewer":
#     st.title("Camera Viewer")

#     # Step 1: Add a file uploader for the zip file
#     uploaded_file = st.file_uploader("Choose a zip file", type="zip")
#     st.write(f"Uploaded file: {uploaded_file}")
#     if uploaded_file is not None:

#         # Step 2: Extract the zip file and read its contents
#         with zipfile.ZipFile(uploaded_file, "r") as zfile:
#             zfile.extractall("temp_folder")

        # # Step 3: Process the image files and the CSV files
        # image_files = []
        # hydrograph_df = None
        # rain_df = None
        # temperature_df = None

        # for root, _, files in os.walk("temp_folder"):
        #     for file in files:
        #         if file.lower().endswith(".jpg"):
        #             image_files.append(os.path.join(root, file))
        #         elif file == "Hydrograph.csv":
        #             hydrograph_df = pd.read_csv(os.path.join(
        #                 root, file), parse_dates=[["Date", "Time"]])
        #         elif file == "Rain.csv":
        #             rain_df = pd.read_csv(os.path.join(
        #                 root, file), parse_dates=[["Date", "Time"]])
        #         elif file == "Temperature.csv":
        #             temperature_df = pd.read_csv(os.path.join(
        #                 root, file), parse_dates=[["Date", "Time"]])

        # st.write(f"Uploaded file: {uploaded_file}")
        # st.write(f"Number of image files: {len(image_files)}")
        # st.write(f"Hydrograph dataframe: {hydrograph_df.head()}")
        # st.write(f"Rain dataframe: {rain_df.head()}")
        # st.write(f"Temperature dataframe: {temperature_df.head()}")

        # # Step 4: For each image file, display the image and plot the graphs
        # for img_file in image_files:
        #     # Extract the time from the image filename
        #     # Remove "image.jpg" from the filename
        #     img_time_str = os.path.basename(img_file)[:-8]
        #     img_time = datetime.strptime(img_time_str, "%m-%d-%Y-%H-%M")

        #     st.image(
        #         img_file, caption=f"Image taken at {img_time.strftime('%Y-%m-%d %H:%M')}")

        #     fig, ax = plt.subplots(3, 1, figsize=(10, 15), sharex=True)
        #     for idx, (df, title, ylabel) in enumerate(zip([hydrograph_df, rain_df, temperature_df], ["Flow", "Rain", "Temperature"], ["Flow", "Rain", "Temperature"])):
        #         # Filter the data within 15 days before and after the image time
        #         mask = (df["Date_Time"] >= img_time - timedelta(days=15)
        #                 ) & (df["Date_Time"] <= img_time + timedelta(days=15))
        #         data = df[mask]

        #         # Plot the graph
        #         ax[idx].plot(data["Date_Time"], data[title], label=title)
        #         ax[idx].set_ylabel(ylabel)

        #         # Add a red dot at the image time
        #         closest_time = data.iloc[(
        #             data["Date_Time"] - img_time).abs().idxmin()]["Date_Time"]
        #         ax[idx].plot(closest_time, data.loc[data["Date_Time"] ==
        #                                             closest_time, title].values[0], "ro", label="Image time")

        #         ax[idx].legend()

        #     plt.xlabel("Date")
        #     plt.xticks(rotation=45)
        #     st.pyplot(fig)

        # # # Clean up the temporary folder
        # # for root, _, files in os.walk("temp_folder"):
        # #     for file in files:
        # #         os.remove(os.path.join(root, file))
        # # os.rmdir("temp_folder")

        # # Clean up the temporary folder
        # shutil.rmtree("temp_folder")

# elif choice == "Camera Viewer":
#     st.title("Camera Viewer")

#     # Step 1: Add a file uploader for the zip file
#     uploaded_file = st.file_uploader("Choose a zip file", type="zip")

#     if uploaded_file is not None:
#         # Step 2: Extract the zip file and read its contents
#         with zipfile.ZipFile(uploaded_file, "r") as zfile:
#             zfile.extractall("temp_folder")

#         # Step 3: Process the image files and the CSV files
#         image_files = []
#         hydrograph_df = None
#         rain_df = None
#         temperature_df = None

#         for root, _, files in os.walk("temp_folder"):
#             for file in files:
#                 if file.endswith(".png"):
#                     image_files.append(os.path.join(root, file))
#                 elif file == "Hydrograph.csv":
#                     hydrograph_df = pd.read_csv(os.path.join(
#                         root, file), parse_dates=[["Date", "Time"]])
#                 elif file == "Rain.csv":
#                     rain_df = pd.read_csv(os.path.join(
#                         root, file), parse_dates=[["Date", "Time"]])
#                 elif file == "Temperature.csv":
#                     temperature_df = pd.read_csv(os.path.join(
#                         root, file), parse_dates=[["Date", "Time"]])

#         # Step 4: For each image file, display the image and plot the graphs
#         for img_file in image_files:
#             # Extract the time from the image filename
#             # Remove "image.png" from the filename
#             img_time_str = os.path.basename(img_file)[:-9]
#             img_time = datetime.strptime(img_time_str, "%m%d%Y%H%M")

#             st.image(
#                 img_file, caption=f"Image taken at {img_time.strftime('%Y-%m-%d %H:%M')}")

#             fig, ax = plt.subplots(3, 1, figsize=(10, 15), sharex=True)
#             for idx, (df, title, ylabel) in enumerate(zip([hydrograph_df, rain_df, temperature_df], ["Flow", "Rain", "Temperature"], ["Flow", "Rain", "Temperature"])):
#                 # Filter the data within 15 days before and after the image time
#                 mask = (df["Date_Time"] >= img_time - timedelta(days=15)
#                         ) & (df["Date_Time"] <= img_time + timedelta(days=15))
#                 data = df[mask]

#                 # Plot the graph
#                 ax[idx].plot(data["Date_Time"], data[title], label=title)
#                 ax[idx].set_ylabel(ylabel)

#                 # Add a red dot at the image time
#                 closest_time = data.iloc[(
#                     data["Date_Time"] - img_time).abs().idxmin()]["Date_Time"]
#                 ax[idx].plot(closest_time, data.loc[data["Date_Time"] ==
#                                                     closest_time, title].values[0], "ro", label="Image time")

#                 ax[idx].legend()

#             plt.xlabel("Date")
#             plt.xticks(rotation=45)
#             st.pyplot(fig)

#         # Clean up the temporary folder
#         for root, _, files in os.walk("temp_folder"):
#             for file in files:
#                 os.remove(os.path.join(root, file))
#         shutil.rmtree("temp_folder")
#         # os.r

# Frequency Analysis page
elif choice == "Frequency Analysis v2":
    st.title('Statistical Distribution Analysis')
    uploaded_file = st.file_uploader("Upload your CSV file", type=["csv"])
    if uploaded_file is not None:
        data = pd.read_csv(uploaded_file, header=None).squeeze()
        if data is not None:
            criteria = []
            criteria.append(fit_and_calculate_criteria(data, stats.norm, 'Normal'))
            # criteria.append(fit_and_calculate_criteria(data, stats.lognorm, 'Log-normal'))
            # criteria.append(fit_and_calculate_criteria(data, stats.genextreme, 'Generalized Extreme Value'))
            # criteria.append(fit_and_calculate_criteria(data, stats.gumbel_r, 'Gumbel'))
            # criteria.append(fit_and_calculate_criteria(data, stats.pearson3, 'Pearson Type 3'))
            st.write(criteria)
            if st.button('Show Normal Plot'):
                create_plot_normal(data)
    else:
        st.stop()
    



    
# Frequency Analysis page
elif choice == "Frequency Analysis":
    st.title('Analyse fréquentielle des débits de crues')

    st.text("Cet outil sélectionne la meilleure distribution pour votre échantillon.")

    uploaded_file = st.file_uploader(
        "Importer un fichier CSV d'une seule colonne qui comprend l'ensemble de l'échantillon.", type="csv")

    if uploaded_file is not None:
        data = pd.read_csv(uploaded_file, header=None, names=['Flow'])
        max_flow = data['Flow'].to_numpy()

        aic_bic_params = {}
        for name, distr in distributions.items():
            aic, bic, params = fit_distribution(distr, max_flow)
            aic_bic_params[name] = {'AIC': aic, 'BIC': bic, 'params': params}

        best_aic_distr = min(
            aic_bic_params, key=lambda x: aic_bic_params[x]['AIC'])
        best_bic_distr = min(
            aic_bic_params, key=lambda x: aic_bic_params[x]['BIC'])

        x = np.linspace(min(max_flow), max(max_flow), 1000)
        for name, info in aic_bic_params.items():
            fig, ax = plt.subplots(figsize=(10, 6))
            ax.hist(max_flow, bins='auto', density=True,
                    alpha=0.6, color='g', label='Histogram')

            params = info['params']
            if name == 'Log-Pearson Type 3':
                ax.plot(x, log_pearson3(x, *params), label=name)
            else:
                distr = distributions[name]
                ax.plot(x, distr.pdf(x, *params), label=name)

            ax.set_xlabel('Flow')
            ax.set_ylabel('Density')
            ax.legend(loc='best')
            plt.savefig(f'{name}_distribution.png', bbox_inches='tight')
            plt.close(fig)

        doc = generate_word_document(
            max_flow, aic_bic_params, best_aic_distr, best_bic_distr)
        st.markdown(download_link(doc, 'Frequency_Analysis.docx'),
                    unsafe_allow_html=True)

    else:
        st.info("Importer votre fichier CSV.")
